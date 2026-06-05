import os
from typing import List, Dict, Any
from pypdf import PdfReader
import docx

class IngestionService:
    @staticmethod
    def parse_pdf(file_path: str) -> List[Dict[str, Any]]:
        """
        Extract pages and text from a PDF document.
        """
        pages = []
        try:
            reader = PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    pages.append({
                        "page": i + 1,
                        "text": text
                    })
        except Exception as e:
            raise ValueError(f"Error parsing PDF file: {str(e)}")
        return pages

    @staticmethod
    def parse_docx(file_path: str) -> List[Dict[str, Any]]:
        """
        Extract text from a Word document.
        Word documents don't have distinct native physical pages in raw python-docx,
        so we chunk paragraphs into logical groups resembling pages (e.g. every 10 paragraphs).
        """
        pages = []
        try:
            doc = docx.Document(file_path)
            current_text = []
            page_idx = 1
            
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    current_text.append(para.text)
                
                # Create a logical page every 12 paragraphs
                if len(current_text) >= 12 or i == len(doc.paragraphs) - 1:
                    pages.append({
                        "page": page_idx,
                        "text": "\n".join(current_text)
                    })
                    current_text = []
                    page_idx += 1
        except Exception as e:
            raise ValueError(f"Error parsing DOCX file: {str(e)}")
        return pages

    @staticmethod
    def parse_txt(file_path: str) -> List[Dict[str, Any]]:
        """
        Extract text from a TXT or Markdown file.
        Chunks text into logical pages based on size (e.g., ~2000 characters).
        """
        pages = []
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            # Split by 2000 characters to form logical pages
            page_size = 2000
            for i in range(0, len(content), page_size):
                pages.append({
                    "page": (i // page_size) + 1,
                    "text": content[i : i + page_size]
                })
        except Exception as e:
            raise ValueError(f"Error parsing text file: {str(e)}")
        return pages

    def parse_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Auto-detects file extension and routes to the appropriate parser.
        """
        _, ext = os.path.splitext(file_path.lower())
        if ext == ".pdf":
            return self.parse_pdf(file_path)
        elif ext == ".docx":
            return self.parse_docx(file_path)
        elif ext in [".txt", ".md", ".markdown"]:
            return self.parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")
